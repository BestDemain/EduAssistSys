# 报告生成服务模块 - 负责生成分析报告

import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from pathlib import Path
import json
import base64
from io import BytesIO
import re
from matplotlib.font_manager import FontProperties
from services.data_service import DataService
from services.analysis_service import AnalysisService

# 设置中文字体
try:
    # 尝试设置中文字体
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
    plt.rcParams['axes.unicode_minus'] = False
except:
    print("警告: 无法设置中文字体，图表中的中文可能无法正确显示")

# 初始化中文字体支持
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# 尝试注册中文字体
try:
    # 注册系统中的中文字体
    pdfmetrics.registerFont(TTFont('SimSun', 'C:\\Windows\\Fonts\\simsun.ttc'))
    pdfmetrics.registerFont(TTFont('SimHei', 'C:\\Windows\\Fonts\\simhei.ttf'))
    # 添加字体映射
    addMapping('SimSun', 0, 0, 'SimSun') # normal
    addMapping('SimHei', 0, 0, 'SimHei') # bold
    print("成功注册中文字体")
except Exception as e:
    print(f"注册中文字体失败: {e}, PDF中的中文可能无法正确显示")

class ReportService:
    def __init__(self):
        self.data_service = DataService()
        self.analysis_service = AnalysisService()
        
        # 报告保存目录
        self.report_dir = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))) / 'reports'
        os.makedirs(self.report_dir, exist_ok=True)
    
    def generate_report(self, report_type='general', student_id=None, format_type='pdf', content=None):
        """生成分析报告
        
        Args:
            report_type: 报告类型，可选值：'general'(总体报告), 'knowledge'(知识点掌握报告), 
                         'behavior'(学习行为报告), 'difficulty'(题目难度报告)
            student_id: 学生ID，如果为None则生成所有学生的汇总报告
            format_type: 报告格式，可选值：'pdf', 'html', 'docx', 'xlsx'
            content: 报告内容配置，指定要包含的内容和图表
            
        Returns:
            生成的报告文件路径
        """
        # 根据报告类型调用相应的报告生成方法
        if report_type == 'knowledge':
            return self._generate_knowledge_report(student_id, format_type, content)
        elif report_type == 'behavior':
            return self._generate_behavior_report(student_id, format_type, content)
        elif report_type == 'difficulty':
            return self._generate_difficulty_report(format_type, content)
        else:  # 默认生成综合报告
            return self._generate_general_report(student_id, format_type, content)
    
    def _generate_knowledge_report(self, student_id=None, format_type='pdf', content=None):
        """生成知识点掌握报告"""
        # 获取知识点掌握度分析结果
        analysis_result = self.analysis_service.analyze_knowledge_mastery(student_id)
        
        if analysis_result['status'] != 'success':
            return {'error': analysis_result['message']}
        
        # 生成报告文件名
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        student_suffix = f"_{student_id}" if student_id else "_all"
        filename = f"knowledge_report{student_suffix}_{timestamp}.{format_type}"
        file_path = self.report_dir / filename
        
        # 根据格式类型生成报告
        if format_type == 'pdf':
            self._generate_pdf_knowledge_report(analysis_result, file_path, content)
        elif format_type == 'html':
            self._generate_html_knowledge_report(analysis_result, file_path, content)
        elif format_type == 'docx':
            self._generate_docx_knowledge_report(analysis_result, file_path, content)
        elif format_type == 'xlsx':
            self._generate_xlsx_knowledge_report(analysis_result, file_path, content)
        else:
            return {'error': f'不支持的报告格式: {format_type}'}
        
        return str(file_path)
    
    def _generate_behavior_report(self, student_id=None, format_type='pdf', content=None):
        """生成学习行为报告"""
        # 获取学习行为分析结果
        analysis_result = self.analysis_service.analyze_learning_behavior(student_id)
        
        if analysis_result['status'] != 'success':
            return {'error': analysis_result['message']}
        
        # 生成报告文件名
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        student_suffix = f"_{student_id}" if student_id else "_all"
        filename = f"behavior_report{student_suffix}_{timestamp}.{format_type}"
        file_path = self.report_dir / filename
        
        # 根据格式类型生成报告
        if format_type == 'pdf':
            self._generate_pdf_behavior_report(analysis_result, file_path, content)
        elif format_type == 'html':
            self._generate_html_behavior_report(analysis_result, file_path, content)
        elif format_type == 'docx':
            self._generate_docx_behavior_report(analysis_result, file_path, content)
        elif format_type == 'xlsx':
            self._generate_xlsx_behavior_report(analysis_result, file_path, content)
        else:
            return {'error': f'不支持的报告格式: {format_type}'}
        
        return str(file_path)
    
    def _generate_difficulty_report(self, format_type='pdf', content=None):
        """生成题目难度报告"""
        # 获取题目难度分析结果
        analysis_result = self.analysis_service.analyze_question_difficulty()
        
        if analysis_result['status'] != 'success':
            return {'error': analysis_result['message']}
        
        # 生成报告文件名
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        filename = f"difficulty_report_{timestamp}.{format_type}"
        file_path = self.report_dir / filename
        
        # 根据格式类型生成报告
        if format_type == 'pdf':
            self._generate_pdf_difficulty_report(analysis_result, file_path, content)
        elif format_type == 'html':
            self._generate_html_difficulty_report(analysis_result, file_path, content)
        elif format_type == 'docx':
            self._generate_docx_difficulty_report(analysis_result, file_path, content)
        elif format_type == 'xlsx':
            self._generate_xlsx_difficulty_report(analysis_result, file_path, content)
        else:
            return {'error': f'不支持的报告格式: {format_type}'}
        
        return str(file_path)
    
    def _generate_general_report(self, student_id=None, format_type='pdf', content=None):
        """生成综合分析报告"""
        # 获取各项分析结果
        knowledge_result = self.analysis_service.analyze_knowledge_mastery(student_id)
        behavior_result = self.analysis_service.analyze_learning_behavior(student_id)
        difficulty_result = self.analysis_service.analyze_question_difficulty()
        
        # 检查分析结果状态
        if knowledge_result['status'] != 'success' or \
           behavior_result['status'] != 'success' or \
           difficulty_result['status'] != 'success':
            return {'error': '获取分析数据失败'}
        
        # 生成报告文件名
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        student_suffix = f"_{student_id}" if student_id else "_all"
        filename = f"general_report{student_suffix}_{timestamp}.{format_type}"
        file_path = self.report_dir / filename
        
        # 合并分析结果
        analysis_result = {
            'knowledge': knowledge_result,
            'behavior': behavior_result,
            'difficulty': difficulty_result
        }
        
        # 根据格式类型生成报告
        if format_type == 'pdf':
            self._generate_pdf_general_report(analysis_result, file_path, content)
        elif format_type == 'html':
            self._generate_html_general_report(analysis_result, file_path, content)
        elif format_type == 'docx':
            self._generate_docx_general_report(analysis_result, file_path, content)
        elif format_type == 'xlsx':
            self._generate_xlsx_general_report(analysis_result, file_path, content)
        else:
            return {'error': f'不支持的报告格式: {format_type}'}
        
        return str(file_path)
    
    # PDF报告生成方法
    def _generate_pdf_knowledge_report(self, analysis_result, file_path, content=None):
        """生成PDF格式的知识点掌握报告"""
        # 这里使用reportlab库生成PDF报告
        # 由于实现复杂，这里仅提供示例框架
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
            from reportlab.lib import colors
            
            # 创建PDF文档
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            styles = getSampleStyleSheet()
            # 修改样式以使用中文字体
            for style_name, style in styles.byName.items():
                if style_name == 'Title':
                    style.fontName = 'SimHei'
                else:
                    style.fontName = 'SimSun'
            story = []
            
            # 添加标题
            title = "知识点掌握度分析报告"
            if analysis_result.get('student_id'):
                title += f" - 学生ID: {analysis_result['student_id']}"
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 12))
            
            # 添加报告生成时间
            story.append(Paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # 添加知识点掌握情况
            story.append(Paragraph("知识点掌握情况", styles['Heading2']))
            story.append(Spacer(1, 6))
            
            # 创建知识点掌握情况表格
            knowledge_data = [['知识点', '正确率', '正确提交率', '平均用时', '总提交次数', '正确提交次数']]
            for knowledge, data in analysis_result['knowledge_mastery'].items():
                knowledge_data.append([
                    knowledge,
                    f"{data['correct_rate']:.2%}",
                    f"{data['correct_submission_rate']:.2%}",
                    f"{data['avg_time_consume']:.2f}秒",
                    str(data['total_submissions']),
                    str(data['correct_submissions'])
                ])
            
            # 创建表格
            table = Table(knowledge_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
            
            # 添加薄弱环节分析
            if analysis_result['weak_points']:
                story.append(Paragraph("薄弱环节分析", styles['Heading2']))
                story.append(Spacer(1, 6))
                
                for weak_point in analysis_result['weak_points']:
                    text = f"知识点 '{weak_point['knowledge']}' "
                    if 'sub_knowledge' in weak_point:
                        text += f"的从属知识点 '{weak_point['sub_knowledge']}' "
                    text += f"正确率为 {weak_point['correct_rate']:.2%}，{weak_point['reason']}。"
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 6))
            
            # 生成PDF
            doc.build(story)
            print(f"已生成PDF报告: {file_path}")
            
        except Exception as e:
            print(f"生成PDF报告失败: {e}")
            return {'error': f'生成PDF报告失败: {e}'}
    
    def _generate_pdf_behavior_report(self, analysis_result, file_path, content=None):
        """生成PDF格式的学习行为报告"""
        # 实现类似于_generate_pdf_knowledge_report的方法
        pass
    
    def _generate_pdf_difficulty_report(self, analysis_result, file_path, content=None):
        """生成PDF格式的题目难度报告"""
        # 实现类似于_generate_pdf_knowledge_report的方法
        pass
    
    def _generate_pdf_general_report(self, analysis_result, file_path, content=None):
        """生成PDF格式的综合分析报告"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
            from reportlab.lib import colors
            
            # 确保文件路径是字符串且目录存在
            file_path = str(file_path)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # 创建PDF文档
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # 修改样式以使用中文字体
            for style_name, style in styles.byName.items():
                if style_name == 'Title':
                    style.fontName = 'SimHei'
                else:
                    style.fontName = 'SimSun'
                    
            story = []
            
            # 添加标题
            title = "综合分析报告"
            if analysis_result.get('knowledge', {}).get('student_id'):
                title += f" - 学生ID: {analysis_result['knowledge']['student_id']}"
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 12))
            
            # 添加报告生成时间
            story.append(Paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # 添加知识点掌握情况
            story.append(Paragraph("知识点掌握情况", styles['Heading2']))
            story.append(Spacer(1, 6))
            
            # 创建知识点掌握情况表格
            knowledge_data = [['知识点', '正确率', '正确提交率', '平均用时', '总提交次数', '正确提交次数']]
            
            knowledge_mastery = analysis_result.get('knowledge', {}).get('knowledge_mastery', {})
            for knowledge, data in knowledge_mastery.items():
                knowledge_data.append([
                    knowledge,
                    f"{data['correct_rate']:.2%}",
                    f"{data['correct_submission_rate']:.2%}",
                    f"{data['avg_time_consume']:.2f}秒",
                    str(data['total_submissions']),
                    str(data['correct_submissions'])
                ])
            
            # 创建表格
            if len(knowledge_data) > 1:  # 确保有数据
                table = Table(knowledge_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'SimHei'),
                    ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
            
            # 添加薄弱环节分析
            weak_points = analysis_result.get('knowledge', {}).get('weak_points', [])
            if weak_points:
                story.append(Paragraph("薄弱环节分析", styles['Heading2']))
                story.append(Spacer(1, 6))
                
                for weak_point in weak_points:
                    text = f"知识点 '{weak_point['knowledge']}' "
                    if 'sub_knowledge' in weak_point:
                        text += f"的从属知识点 '{weak_point['sub_knowledge']}' "
                    text += f"正确率为 {weak_point['correct_rate']:.2%}，{weak_point['reason']}。"
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 6))
            
            # 添加学习行为分析
            story.append(Paragraph("学习行为分析", styles['Heading2']))
            story.append(Spacer(1, 6))
            
            # 获取学习行为数据
            behavior_data = analysis_result.get('behavior', {}).get('behavior_profile', {})
            if behavior_data:
                # 添加答题时间分布
                if 'peak_hours' in behavior_data:
                    story.append(Paragraph("答题时间分布", styles['Heading3']))
                    story.append(Spacer(1, 4))
                    peak_data = [["时间段", "提交次数"]]
                    for hour_data in behavior_data['peak_hours']:
                        hour = hour_data['hour']
                        count = hour_data['count']
                        peak_data.append([f"{hour}:00-{int(hour)+1}:00", str(count)])
                    
                    peak_table = Table(peak_data)
                    peak_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'SimHei'),
                        ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(peak_table)
                    story.append(Spacer(1, 8))
                
                # 添加答题状态分布
                if 'state_distribution' in behavior_data:
                    story.append(Paragraph("答题状态分布", styles['Heading3']))
                    story.append(Spacer(1, 4))
                    state_data = [["提交状态", "次数", "百分比"]]
                    for state, count in behavior_data['state_distribution'].items():
                        total = sum(behavior_data['state_distribution'].values())
                        percentage = count / total if total > 0 else 0
                        state_data.append([state, str(count), f"{percentage:.2%}"])
                    
                    state_table = Table(state_data)
                    state_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'SimHei'),
                        ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(state_table)
                    story.append(Spacer(1, 8))
                
                # 添加学习行为模式分析
                story.append(Paragraph("学习行为模式分析", styles['Heading3']))
                story.append(Spacer(1, 4))
                
                # 创建学习行为模式文本
                behavior_patterns = []
                
                # 根据正确率和用时分析学习效率
                if 'correct_rate' in behavior_data:
                    correct_rate = behavior_data['correct_rate']
                    if correct_rate > 0.8:
                        behavior_patterns.append("学生答题正确率很高，掌握知识点较好。")
                    elif correct_rate > 0.6:
                        behavior_patterns.append("学生答题正确率一般，对部分知识点的理解需要加强。")
                    else:
                        behavior_patterns.append("学生答题正确率较低，需要加强基础知识学习。")
                
                # 根据答题时间分析学习时间习惯
                if 'peak_hours' in behavior_data and behavior_data['peak_hours']:
                    peak_hour = behavior_data['peak_hours'][0]['hour']
                    if 8 <= peak_hour <= 12:
                        behavior_patterns.append("学生倾向于在上午时间段答题，精力较为充沛。")
                    elif 13 <= peak_hour <= 17:
                        behavior_patterns.append("学生倾向于在下午时间段答题，学习效率较为稳定。")
                    elif 18 <= peak_hour <= 22:
                        behavior_patterns.append("学生倾向于在晚上时间段答题，可能会影响休息时间。")
                    else:
                        behavior_patterns.append(f"学生在{peak_hour}:00-{peak_hour+1}:00时段答题最多，可能是非常规学习时间。")
                
                # 添加方法使用分析
                if 'method_distribution' in behavior_data:
                    methods = behavior_data['method_distribution']
                    if methods and len(methods) > 0:
                        most_method = max(methods.items(), key=lambda x: x[1])[0]
                        behavior_patterns.append(f"学生最常使用的答题方法是{most_method}。")
                
                # 如果有相对表现数据，添加相对表现分析
                if 'relative_performance' in behavior_data:
                    rel_perf = behavior_data['relative_performance']
                    if 'correct_rate_vs_avg' in rel_perf:
                        rate_vs_avg = rel_perf['correct_rate_vs_avg']
                        if rate_vs_avg > 0.1:
                            behavior_patterns.append("学生的正确率明显高于平均水平，表现优秀。")
                        elif rate_vs_avg < -0.1:
                            behavior_patterns.append("学生的正确率低于平均水平，需要改进。")
                
                # 添加行为模式文本
                if behavior_patterns:
                    for pattern in behavior_patterns:
                        story.append(Paragraph(f"- {pattern}", styles['Normal']))
                        story.append(Spacer(1, 4))
                else:
                    story.append(Paragraph("无法生成行为模式分析，数据不足。", styles['Normal']))
                    story.append(Spacer(1, 4))
            else:
                story.append(Paragraph("无学习行为分析数据", styles['Normal']))
            
            story.append(Spacer(1, 12))
            
            # 添加题目难度分析
            story.append(Paragraph("题目难度分析", styles['Heading2']))
            story.append(Spacer(1, 6))
            
            # 获取题目难度数据
            difficulty_data = analysis_result.get('difficulty', {})
            
            # 计算题目难度分布
            if 'question_difficulty' in difficulty_data:
                # 从question_difficulty中计算难度分布
                difficulty_distribution = {}
                for title_id, question in difficulty_data['question_difficulty'].items():
                    # 根据正确率确定难度级别
                    correct_rate = question['correct_rate']
                    if correct_rate >= 0.8:
                        level = '容易'
                    elif correct_rate >= 0.5:
                        level = '中等'
                    else:
                        level = '困难'
                    
                    difficulty_distribution[level] = difficulty_distribution.get(level, 0) + 1
                
                # 添加题目难度分布
                if difficulty_distribution:
                    story.append(Paragraph("题目难度分布", styles['Heading3']))
                    story.append(Spacer(1, 4))
                    diff_data = [["难度级别", "题目数量", "百分比"]]
                    for level, count in difficulty_distribution.items():
                        total = sum(difficulty_distribution.values())
                        percentage = count / total if total > 0 else 0
                        diff_data.append([level, str(count), f"{percentage:.2%}"])
                    
                    diff_table = Table(diff_data)
                    diff_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'SimHei'),
                        ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(diff_table)
                    story.append(Spacer(1, 8))
                
                # 添加难度异常题目分析
                if 'unreasonable_questions' in difficulty_data:
                    story.append(Paragraph("难度异常题目分析", styles['Heading3']))
                    story.append(Spacer(1, 4))
                    abnormal_data = [["题目ID", "正确率", "平均掌握度", "知识点", "原因"]]
                    for question in difficulty_data['unreasonable_questions']:
                        abnormal_data.append([
                            question['title_id'],
                            f"{question['correct_rate']:.2%}",
                            f"{question['avg_mastery']:.2%}",
                            question['knowledge'],
                            question['reason']
                        ])
                    
                    if len(abnormal_data) > 1:  # 确保有数据
                        abnormal_table = Table(abnormal_data)
                        abnormal_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'SimHei'),
                            ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        story.append(abnormal_table)
                    else:
                        story.append(Paragraph("无难度异常题目", styles['Normal']))
                    story.append(Spacer(1, 8))
                
                # 添加难度建议
                story.append(Paragraph("难度建议", styles['Heading3']))
                story.append(Spacer(1, 4))
                
                # 生成难度建议
                difficulty_suggestions = []
                
                # 根据难度分布生成建议
                if difficulty_distribution:
                    # 分析难度分布情况
                    total_questions = sum(difficulty_distribution.values())
                    difficult_count = difficulty_distribution.get('困难', 0)
                    easy_count = difficulty_distribution.get('容易', 0)
                    
                    # 如果困难题目比例过高
                    if difficult_count / total_questions > 0.4:
                        difficulty_suggestions.append("困难题目比例较高，建议适当增加中等难度和简单题目，以提高学生学习积极性。")
                    
                    # 如果简单题目比例过高
                    if easy_count / total_questions > 0.6:
                        difficulty_suggestions.append("简单题目比例较高，建议适当增加难度，以提升学生解决问题的能力。")
                    
                    # 如果难度分布较为均衡
                    if 0.2 <= difficult_count / total_questions <= 0.4 and 0.3 <= easy_count / total_questions <= 0.5:
                        difficulty_suggestions.append("题目难度分布较为均衡，适合大多数学生的学习需求。")
                
                # 根据异常题目生成建议
                if 'unreasonable_questions' in difficulty_data and difficulty_data['unreasonable_questions']:
                    unreasonable_count = len(difficulty_data['unreasonable_questions'])
                    difficulty_suggestions.append(f"有{unreasonable_count}道题目的难度与学生知识掌握程度不符，建议检查这些题目的设计。")
                    
                    # 如果异常题目较多，给出进一步建议
                    if unreasonable_count >= 3:
                        difficulty_suggestions.append("建议对难度异常题目进行重新设计或提供更详细的解题指导。")
                
                # 添加建议到报告中
                if difficulty_suggestions:
                    for suggestion in difficulty_suggestions:
                        story.append(Paragraph(f"- {suggestion}", styles['Normal']))
                        story.append(Spacer(1, 4))
                else:
                    story.append(Paragraph("无难度建议", styles['Normal']))
            else:
                story.append(Paragraph("无题目难度分析数据", styles['Normal']))
            
            # 生成PDF
            doc.build(story)
            print(f"已生成PDF报告: {file_path}")
            
        except Exception as e:
            print(f"生成PDF报告失败: {e}")
            return {'error': f'生成PDF报告失败: {e}'}
    


    # HTML报告生成方法
    # 提取公共方法
    def _generate_html_header(self, title, student_id=None):
        """生成HTML头部公共部分"""
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #333; }}
                h2 {{ color: #666; }}
                table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                tr:nth-child(even) {{ background-color: #f9f9f9; }}
                .weak-point {{ color: #d9534f; }}
                .section {{ margin-bottom: 30px; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            {f'<p>学生ID: {student_id}</p>' if student_id else ''}
            <p>报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        """

    def _generate_knowledge_section(self, knowledge_data):
        """生成知识点分析公共部分，包含可视化图表"""
        if not knowledge_data.get('knowledge_mastery'):
            return "<p>无知识点掌握数据</p>"
        
        # 准备图表数据
        knowledge_names = []
        correct_rates = []
        submission_rates = []
        total_submissions = []
        correct_submissions = []
        
        for knowledge, data in knowledge_data['knowledge_mastery'].items():
            knowledge_names.append(knowledge)
            correct_rates.append(data['correct_rate'] * 100)  # 转换为百分比
            submission_rates.append(data['correct_submission_rate'] * 100)
            total_submissions.append(data['total_submissions'])
            correct_submissions.append(data['correct_submissions'])
        
        html = """
        <div class='section'>
            <h2>知识点掌握情况</h2>
            
            <!-- 正确率柱状图 -->
            <div class='chart-container'>
                <canvas id='correctRateChart'></canvas>
            </div>
            
            <!-- 提交次数折线图 -->
            <div class='chart-container'>
                <canvas id='submissionChart'></canvas>
            </div>
            
            <table>
                <tr>
                    <th>知识点</th>
                    <th>正确率</th>
                    <th>正确提交率</th>
                    <th>平均用时</th>
                    <th>总提交次数</th>
                    <th>正确提交次数</th>
                </tr>
        """
        
        for knowledge, data in knowledge_data['knowledge_mastery'].items():
            html += f"""
                <tr>
                    <td>{knowledge}</td>
                    <td>{data['correct_rate']:.2%}</td>
                    <td>{data['correct_submission_rate']:.2%}</td>
                    <td>{data['avg_time_consume']:.2f}秒</td>
                    <td>{data['total_submissions']}</td>
                    <td>{data['correct_submissions']}</td>
                </tr>
            """
        
        html += "</table>"
        
        if knowledge_data.get('weak_points'):
            html += "<h3>薄弱环节分析</h3><ul>"
            for weak_point in knowledge_data['weak_points']:
                text = f"知识点 '{weak_point['knowledge']}' "
                if 'sub_knowledge' in weak_point:
                    text += f"的从属知识点 '{weak_point['sub_knowledge']}' "
                text += f"正确率为 {weak_point['correct_rate']:.2%}，{weak_point['reason']}。"
                html += f"<li class='weak-point'>{text}</li>"
            html += "</ul>"
        
        # 添加图表JS代码
        html += """
        <script src='https://cdn.jsdelivr.net/npm/chart.js'></script>
        <script>
            // 正确率柱状图
            const correctRateCtx = document.getElementById('correctRateChart').getContext('2d');
            new Chart(correctRateCtx, {
                type: 'bar',
                data: {
                    labels: """ + str(knowledge_names) + """,
                    datasets: [{
                        label: '知识点正确率(%)',
                        data: """ + str(correct_rates) + """,
                        backgroundColor: 'rgba(54, 162, 235, 0.5)',
                        borderColor: 'rgba(54, 162, 235, 1)',
                        borderWidth: 1
                    }]
                },
                options: {
                    responsive: true,
                    scales: {
                        y: {
                            beginAtZero: true,
                            max: 100,
                            title: {
                                display: true,
                                text: '正确率(%)'
                            }
                        }
                    }
                }
            });
            
            // 提交次数折线图
            const submissionCtx = document.getElementById('submissionChart').getContext('2d');
            new Chart(submissionCtx, {
                type: 'line',
                data: {
                    labels: """ + str(knowledge_names) + """,
                    datasets: [
                        {
                            label: '总提交次数',
                            data: """ + str(total_submissions) + """,
                            borderColor: 'rgba(255, 99, 132, 1)',
                            backgroundColor: 'rgba(255, 99, 132, 0.1)',
                            borderWidth: 2,
                            tension: 0.1
                        },
                        {
                            label: '正确提交次数',
                            data: """ + str(correct_submissions) + """,
                            borderColor: 'rgba(75, 192, 192, 1)',
                            backgroundColor: 'rgba(75, 192, 192, 0.1)',
                            borderWidth: 2,
                            tension: 0.1
                        }
                    ]
                },
                options: {
                    responsive: true,
                    scales: {
                        y: {
                            beginAtZero: true,
                            title: {
                                display: true,
                                text: '提交次数'
                            }
                        }
                    }
                }
            });
        </script>
        
        <style>
            .chart-container {
                position: relative;
                height: 300px;
                margin-bottom: 30px;
            }
            .weak-point {
                color: red;
                margin-bottom: 8px;
            }
        </style>
        </div>
        """
        return html

    # def _generate_html_knowledge_report(self, analysis_result, file_path, content=None):
    #     """生成HTML格式的知识点掌握报告"""
    #     # 使用简单的HTML模板生成报告
    #     try:
    #         # 创建HTML内容
    #         html_content = f"""
    #         <!DOCTYPE html>
    #         <html>
    #         <head>
    #             <meta charset="UTF-8">
    #             <title>知识点掌握度分析报告</title>
    #             <style>
    #                 body {{ font-family: Arial, sans-serif; margin: 20px; }}
    #                 h1 {{ color: #333; }}
    #                 h2 {{ color: #666; }}
    #                 table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
    #                 th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
    #                 th {{ background-color: #f2f2f2; }}
    #                 tr:nth-child(even) {{ background-color: #f9f9f9; }}
    #                 .weak-point {{ color: #d9534f; }}
    #             </style>
    #         </head>
    #         <body>
    #             <h1>知识点掌握度分析报告</h1>
    #         """
            
    #         # 添加学生信息（如果有）
    #         if analysis_result.get('student_id'):
    #             html_content += f"<p>学生ID: {analysis_result['student_id']}</p>"
            
    #         # 添加报告生成时间
    #         html_content += f"<p>报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>"
            
    #         # 添加知识点掌握情况表格
    #         html_content += """
    #             <h2>知识点掌握情况</h2>
    #             <table>
    #                 <tr>
    #                     <th>知识点</th>
    #                     <th>正确率</th>
    #                     <th>正确提交率</th>
    #                     <th>平均用时</th>
    #                     <th>总提交次数</th>
    #                     <th>正确提交次数</th>
    #                 </tr>
    #         """
            
    #         for knowledge, data in analysis_result['knowledge_mastery'].items():
    #             html_content += f"""
    #                 <tr>
    #                     <td>{knowledge}</td>
    #                     <td>{data['correct_rate']:.2%}</td>
    #                     <td>{data['correct_submission_rate']:.2%}</td>
    #                     <td>{data['avg_time_consume']:.2f}秒</td>
    #                     <td>{data['total_submissions']}</td>
    #                     <td>{data['correct_submissions']}</td>
    #                 </tr>
    #             """
            
    #         html_content += "</table>"
            
    #         # 添加薄弱环节分析
    #         if analysis_result['weak_points']:
    #             html_content += "<h2>薄弱环节分析</h2><ul>"
                
    #             for weak_point in analysis_result['weak_points']:
    #                 text = f"知识点 '{weak_point['knowledge']}' "
    #                 if 'sub_knowledge' in weak_point:
    #                     text += f"的从属知识点 '{weak_point['sub_knowledge']}' "
    #                 text += f"正确率为 {weak_point['correct_rate']:.2%}，{weak_point['reason']}。"
    #                 html_content += f"<li class='weak-point'>{text}</li>"
                
    #             html_content += "</ul>"
            
    #         # 结束HTML
    #         html_content += """
    #         </body>
    #         </html>
    #         """
            
    #         # 写入文件
    #         with open(file_path, 'w', encoding='utf-8') as f:
    #             f.write(html_content)
            
    #         print(f"已生成HTML报告: {file_path}")
            
    #     except Exception as e:
    #         print(f"生成HTML报告失败: {e}")
    #         return {'error': f'生成HTML报告失败: {e}'}

    # 有复用
    def _generate_html_knowledge_report(self, analysis_result, file_path, content=None):
        try:
            html = self._generate_html_header(
                "知识点掌握度分析报告", 
                analysis_result.get('student_id')
            )
            
            html += self._generate_knowledge_section(analysis_result)
            
            html += """
            </body>
            </html>
            """
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html)
                
        except Exception as e:
            print(f"生成HTML报告失败: {e}")
            return {'error': f'生成HTML报告失败: {e}'}
    


    # def _generate_html_behavior_report(self, analysis_result, file_path, content=None):
    #     """生成HTML格式的学习行为报告"""
    #     # 实现类似于_generate_html_knowledge_report的方法
    #     pass
    
    # def _generate_html_difficulty_report(self, analysis_result, file_path, content=None):
    #     """生成HTML格式的题目难度报告"""
    #     # 实现类似于_generate_html_knowledge_report的方法
    #     pass
    
    
    def _generate_html_behavior_report(self, analysis_result, file_path, content=None):
        """生成HTML格式的学习行为报告（复用版）"""
        try:
            # 获取学生ID（如果存在）
            student_id = analysis_result.get('student_id')
            
            # 构建HTML
            html = self._generate_html_header("学习行为分析报告", student_id)
            
            # 添加行为分析部分
            behavior_data = analysis_result.get('behavior', {})
            html += self._generate_behavior_section(behavior_data)
            
            # 结束HTML
            html += """
            </body>
            </html>
            """
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            print(f"已生成行为分析HTML报告: {file_path}")
            
        except Exception as e:
            print(f"生成行为分析报告失败: {e}")
            return {'error': f'生成行为分析报告失败: {e}'}

    def _generate_html_difficulty_report(self, analysis_result, file_path, content=None):
        """生成HTML格式的题目难度报告（复用版）"""
        try:
            # 获取学生ID（如果存在）
            student_id = analysis_result.get('student_id')
            
            # 构建HTML
            html = self._generate_html_header("题目难度分析报告", student_id)
            
            # 添加难度分析部分
            difficulty_data = analysis_result.get('difficulty', {})
            html += self._generate_difficulty_section(difficulty_data)
            
            # 结束HTML
            html += """
            </body>
            </html>
            """
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            print(f"已生成难度分析HTML报告: {file_path}")
            
        except Exception as e:
            print(f"生成难度分析报告失败: {e}")
            return {'error': f'生成难度分析报告失败: {e}'}





    def _generate_behavior_section(self, behavior_data):
        """生成行为分析公共部分，包含可视化图表"""
        if not behavior_data.get('behavior_profile'):
            return "<p>无学习行为分析数据</p>"
        
        behavior_profile = behavior_data['behavior_profile']
        html = """
        <div class='section'>
            <h2>学习行为分析</h2>
            
            <!-- 答题时间分布图表容器 -->
            <div class='chart-container'>
                <canvas id='timeDistributionChart'></canvas>
            </div>
            
            <!-- 答题状态分布图表容器 -->
            <div class='chart-container'>
                <canvas id='stateDistributionChart'></canvas>
            </div>
        """
        
        # 准备图表数据
        time_labels = []
        time_counts = []
        state_labels = []
        state_counts = []
        state_percentages = []
        
        # 答题时间分布
        if 'peak_hours' in behavior_profile:
            html += "<h3>答题时间分布</h3>"
            html += """
                <table>
                    <tr>
                        <th>时间段</th>
                        <th>提交次数</th>
                    </tr>
            """
            
            for item in behavior_profile['peak_hours']:
                hour = item['hour']
                count = item['count']
                time_labels.append(f"{hour}:00-{int(hour)+1}:00")
                time_counts.append(count)
                html += f"""
                    <tr>
                        <td>{hour}:00-{int(hour)+1}:00</td>
                        <td>{count}</td>
                    </tr>
                """
            
            html += "</table>"
        
        # 答题状态分布
        if 'state_distribution' in behavior_profile:
            html += "<h3>答题状态分布</h3>"
            html += """
                <table>
                    <tr>
                        <th>提交状态</th>
                        <th>次数</th>
                        <th>百分比</th>
                    </tr>
            """
            
            total = sum(behavior_profile['state_distribution'].values())
            for state, count in behavior_profile['state_distribution'].items():
                percentage = count / total if total > 0 else 0
                state_labels.append(state)
                state_counts.append(count)
                state_percentages.append(percentage)
                html += f"""
                    <tr>
                        <td>{state}</td>
                        <td>{count}</td>
                        <td>{percentage:.2%}</td>
                    </tr>
                """
            
            html += "</table>"
        
        # 行为模式分析
        html += "<h3>学习行为模式分析</h3><ul>"
        patterns = self._generate_behavior_patterns(behavior_profile)
        
        if patterns:
            for pattern in patterns:
                html += f"<li>{pattern}</li>"
        else:
            html += "<li>无法生成行为模式分析，数据不足。</li>"
        
        # 添加图表JS代码
        html += """
        <script src='https://cdn.jsdelivr.net/npm/chart.js'></script>
        <script>
            // 答题时间分布柱状图
            const timeCtx = document.getElementById('timeDistributionChart').getContext('2d');
            new Chart(timeCtx, {
                type: 'bar',
                data: {
                    labels: """ + str(time_labels) + """,
                    datasets: [{
                        label: '提交次数',
                        data: """ + str(time_counts) + """,
                        backgroundColor: 'rgba(54, 162, 235, 0.7)',
                        borderColor: 'rgba(54, 162, 235, 1)',
                        borderWidth: 1
                    }]
                },
                options: {
                    responsive: true,
                    plugins: {
                        title: {
                            display: true,
                            text: '答题时间分布',
                            font: {
                                size: 16
                            }
                        }
                    },
                    scales: {
                        y: {
                            beginAtZero: true,
                            title: {
                                display: true,
                                text: '提交次数'
                            }
                        },
                        x: {
                            title: {
                                display: true,
                                text: '时间段'
                            }
                        }
                    }
                }
            });
            
            // 答题状态分布饼图
            const stateCtx = document.getElementById('stateDistributionChart').getContext('2d');
            new Chart(stateCtx, {
                type: 'pie',
                data: {
                    labels: """ + str(state_labels) + """,
                    datasets: [{
                        data: """ + str(state_counts) + """,
                        backgroundColor: [
                            'rgba(255, 99, 132, 0.7)',
                            'rgba(54, 162, 235, 0.7)',
                            'rgba(255, 206, 86, 0.7)',
                            'rgba(75, 192, 192, 0.7)',
                            'rgba(153, 102, 255, 0.7)',
                            'rgba(255, 159, 64, 0.7)',
                            'rgba(199, 199, 199, 0.7)',
                            'rgba(83, 102, 255, 0.7)',
                            'rgba(40, 159, 64, 0.7)',
                            'rgba(210, 99, 132, 0.7)',
                            'rgba(100, 162, 235, 0.7)'
                        ],
                        borderColor: [
                            'rgba(255, 99, 132, 1)',
                            'rgba(54, 162, 235, 1)',
                            'rgba(255, 206, 86, 1)',
                            'rgba(75, 192, 192, 1)',
                            'rgba(153, 102, 255, 1)',
                            'rgba(255, 159, 64, 1)',
                            'rgba(199, 199, 199, 1)',
                            'rgba(83, 102, 255, 1)',
                            'rgba(40, 159, 64, 1)',
                            'rgba(210, 99, 132, 1)',
                            'rgba(100, 162, 235, 1)'
                        ],
                        borderWidth: 1
                    }]
                },
                options: {
                    responsive: true,
                    plugins: {
                        title: {
                            display: true,
                            text: '答题状态分布',
                            font: {
                                size: 16
                            }
                        },
                        tooltip: {
                            callbacks: {
                                label: function(context) {
                                    const label = context.label || '';
                                    const value = context.raw || 0;
                                    const total = context.dataset.data.reduce((a, b) => a + b, 0);
                                    const percentage = Math.round((value / total) * 100);
                                    return `${label}: ${value} (${percentage}%)`;
                                }
                            }
                        },
                        legend: {
                            position: 'right',
                            labels: {
                                boxWidth: 12
                            }
                        }
                    }
                }
            });
        </script>
        
        <style>
            .chart-container {
                position: relative;
                height: 400px;
                margin-bottom: 40px;
            }
            .section table {
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 20px;
            }
            .section th, .section td {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            .section th {
                background-color: #f2f2f2;
            }
            .section ul {
                padding-left: 20px;
            }
        </style>
        </div>
        """
        return html

    def _generate_behavior_patterns(self, behavior_profile):
        """生成行为模式分析文本"""
        patterns = []
        
        # 根据正确率分析学习效率
        if 'correct_rate' in behavior_profile:
            correct_rate = behavior_profile['correct_rate']
            if correct_rate > 0.8:
                patterns.append("学生答题正确率很高，掌握知识点较好。")
            elif correct_rate > 0.6:
                patterns.append("学生答题正确率一般，对部分知识点的理解需要加强。")
            else:
                patterns.append("学生答题正确率较低，需要加强基础知识学习。")
        
        # 根据答题时间分析学习时间习惯
        if 'peak_hours' in behavior_profile and behavior_profile['peak_hours']:
            peak_hour = behavior_profile['peak_hours'][0]['hour']
            if 8 <= peak_hour <= 12:
                patterns.append("学生倾向于在上午时间段答题，精力较为充沛。")
            elif 13 <= peak_hour <= 17:
                patterns.append("学生倾向于在下午时间段答题，学习效率较为稳定。")
            elif 18 <= peak_hour <= 22:
                patterns.append("学生倾向于在晚上时间段答题，可能会影响休息时间。")
            else:
                patterns.append(f"学生在{peak_hour}:00-{peak_hour+1}:00时段答题最多，可能是非常规学习时间。")
        
        # 方法使用分析
        if 'method_distribution' in behavior_profile:
            methods = behavior_profile['method_distribution']
            if methods and len(methods) > 0:
                most_method = max(methods.items(), key=lambda x: x[1])[0]
                patterns.append(f"学生最常使用的答题方法是{most_method}。")
        
        # 相对表现分析
        if 'relative_performance' in behavior_profile:
            rel_perf = behavior_profile['relative_performance']
            if 'correct_rate_vs_avg' in rel_perf:
                rate_vs_avg = rel_perf['correct_rate_vs_avg']
                if rate_vs_avg > 0.1:
                    patterns.append("学生的正确率明显高于平均水平，表现优秀。")
                elif rate_vs_avg < -0.1:
                    patterns.append("学生的正确率低于平均水平，需要改进。")
        
        return patterns

    def _generate_difficulty_section(self, difficulty_data):
        """生成难度分析公共部分"""
        if not difficulty_data.get('question_difficulty'):
            return "<p>无题目难度分析数据</p>"
        
        html = """
        <div class='section'>
            <h2>题目难度分析</h2>
            <!-- 添加Chart.js库 -->
            <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
        """
        
        question_difficulty = difficulty_data['question_difficulty']
        
        # 难度分布
        difficulty_levels = {'简单': 0, '中等': 0, '困难': 0}
        
        for title_id, data in question_difficulty.items():
            if data['correct_rate'] > 0.7:
                difficulty_levels['简单'] += 1
            elif data['correct_rate'] > 0.4:
                difficulty_levels['中等'] += 1
            else:
                difficulty_levels['困难'] += 1
        
        total = sum(difficulty_levels.values())
        
        # 难度分布饼图
        html += """
        <div style="display: flex; justify-content: space-between; margin-bottom: 20px;">
            <div style="width: 45%;">
                <h3>题目难度分布</h3>
                <canvas id="difficultyChart"></canvas>
            </div>
        """
        
        # 异常题目分析
        unreasonable_questions = difficulty_data.get('unreasonable_questions', [])
        
        # 异常题目正确率分布柱状图
        if unreasonable_questions:
            html += """
            <div style="width: 45%;">
                <h3>异常题目正确率分布</h3>
                <canvas id="unreasonableChart"></canvas>
            </div>
            """
        
        html += "</div>"
        
        # 难度分布表格
        html += """
            <table>
                <tr>
                    <th>难度级别</th>
                    <th>题目数量</th>
                    <th>百分比</th>
                </tr>
        """
        
        for level, count in difficulty_levels.items():
            percentage = count / total if total > 0 else 0
            html += f"""
                <tr>
                    <td>{level}</td>
                    <td>{count}</td>
                    <td>{percentage:.2%}</td>
                </tr>
            """
        
        html += "</table>"
        
        # 异常题目分析表格
        if unreasonable_questions:
            html += "<h3>难度异常题目分析</h3>"
            html += """
                <table>
                    <tr>
                        <th>题目ID</th>
                        <th>正确率</th>
                        <th>平均掌握度</th>
                        <th>知识点</th>
                        <th>原因</th>
                    </tr>
            """
            
            for question in unreasonable_questions:
                html += f"""
                    <tr>
                        <td>{question['title_id']}</td>
                        <td>{question['correct_rate']:.2%}</td>
                        <td>{question['avg_mastery']:.2%}</td>
                        <td>{question['knowledge']}</td>
                        <td>{question['reason']}</td>
                    </tr>
                """
            
            html += "</table>"
        else:
            html += "<p>无难度异常题目</p>"
        
        # 难度建议
        html += "<h3>难度建议</h3><ul>"
        suggestions = self._generate_difficulty_suggestions(difficulty_levels, unreasonable_questions)
        
        if suggestions:
            for suggestion in suggestions:
                html += f"<li>{suggestion}</li>"
        else:
            html += "<li>无难度建议</li>"
        
        html += "</ul>"
        
        # 添加图表渲染脚本
        html += """
        <script>
            // 难度分布饼图
            const difficultyCtx = document.getElementById('difficultyChart');
            new Chart(difficultyCtx, {
                type: 'pie',
                data: {
                    labels: ['简单', '中等', '困难'],
                    datasets: [{
                        data: ["""
        html += f"{difficulty_levels['简单']}, {difficulty_levels['中等']}, {difficulty_levels['困难']}"
        html += """],
                        backgroundColor: [
                            'rgba(75, 192, 192, 0.7)',
                            'rgba(255, 206, 86, 0.7)',
                            'rgba(255, 99, 132, 0.7)'
                        ],
                        borderColor: [
                            'rgba(75, 192, 192, 1)',
                            'rgba(255, 206, 86, 1)',
                            'rgba(255, 99, 132, 1)'
                        ],
                        borderWidth: 1
                    }]
                },
                options: {
                    responsive: true,
                    plugins: {
                        legend: {
                            position: 'top',
                        },
                        tooltip: {
                            callbacks: {
                                label: function(context) {
                                    const label = context.label || '';
                                    const value = context.raw || 0;
                                    const total = context.dataset.data.reduce((a, b) => a + b, 0);
                                    const percentage = Math.round((value / total) * 100);
                                    return `${label}: ${value} (${percentage}%)`;
                                }
                            }
                        }
                    }
                }
            });
        """
        
        if unreasonable_questions:
            # 准备异常题目正确率数据
            correct_rates = [q['correct_rate'] for q in unreasonable_questions]
            knowledge_points = list(set(q['knowledge'] for q in unreasonable_questions))
            avg_rates_by_knowledge = []
            
            for knowledge in knowledge_points:
                rates = [q['correct_rate'] for q in unreasonable_questions if q['knowledge'] == knowledge]
                avg_rate = sum(rates) / len(rates) if rates else 0
                avg_rates_by_knowledge.append(avg_rate)
            
            html += """
                // 异常题目正确率柱状图
                const unreasonableCtx = document.getElementById('unreasonableChart');
                new Chart(unreasonableCtx, {
                    type: 'bar',
                    data: {
                        labels: """
            html += str(knowledge_points).replace("'", '"')
            html += """,
                        datasets: [{
                            label: '平均正确率',
                            data: """
            html += str(avg_rates_by_knowledge)
            html += """,
                            backgroundColor: 'rgba(54, 162, 235, 0.7)',
                            borderColor: 'rgba(54, 162, 235, 1)',
                            borderWidth: 1
                        }]
                    },
                    options: {
                        responsive: true,
                        scales: {
                            y: {
                                beginAtZero: true,
                                ticks: {
                                    callback: function(value) {
                                        return (value * 100).toFixed(0) + '%';
                                    }
                                }
                            }
                        },
                        plugins: {
                            tooltip: {
                                callbacks: {
                                    label: function(context) {
                                        return '平均正确率: ' + (context.raw * 100).toFixed(2) + '%';
                                    }
                                }
                            }
                        }
                    }
                });
            """
        
        html += """
        </script>
        </div>
        """
        return html

    def _generate_difficulty_suggestions(self, difficulty_levels, unreasonable_questions):
        """生成难度建议文本"""
        suggestions = []
        total_questions = sum(difficulty_levels.values())
        
        if total_questions == 0:
            return suggestions
        
        difficult_count = difficulty_levels.get('困难', 0)
        easy_count = difficulty_levels.get('简单', 0)
        
        # 难度分布建议
        if difficult_count / total_questions > 0.4:
            suggestions.append("困难题目比例较高，建议适当增加中等难度和简单题目，以提高学生学习积极性。")
        
        if easy_count / total_questions > 0.6:
            suggestions.append("简单题目比例较高，建议适当增加难度，以提升学生解决问题的能力。")
        
        if (0.2 <= difficult_count / total_questions <= 0.4 and 
            0.3 <= easy_count / total_questions <= 0.5):
            suggestions.append("题目难度分布较为均衡，适合大多数学生的学习需求。")
        
        # 异常题目建议
        if unreasonable_questions:
            unreasonable_count = len(unreasonable_questions)
            suggestions.append(f"有{unreasonable_count}道题目的难度与学生知识掌握程度不符，建议检查这些题目的设计。")
            
            if unreasonable_count >= 3:
                suggestions.append("建议对难度异常题目进行重新设计或提供更详细的解题指导。")
        
        return suggestions




    # def _generate_html_general_report(self, analysis_result, file_path, content=None):
    #     """生成HTML格式的综合分析报告"""
    #     try:
    #         # 创建HTML内容
    #         html_content = f"""
    #         <!DOCTYPE html>
    #         <html>
    #         <head>
    #             <meta charset="UTF-8">
    #             <title>综合分析报告</title>
    #             <style>
    #                 body {{ font-family: Arial, sans-serif; margin: 20px; }}
    #                 h1 {{ color: #333; }}
    #                 h2 {{ color: #666; }}
    #                 h3 {{ color: #888; }}
    #                 table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
    #                 th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
    #                 th {{ background-color: #f2f2f2; }}
    #                 tr:nth-child(even) {{ background-color: #f9f9f9; }}
    #                 .weak-point {{ color: #d9534f; }}
    #                 .section {{ margin-bottom: 30px; }}
    #             </style>
    #         </head>
    #         <body>
    #             <h1>综合分析报告</h1>
    #         """
            
    #         # 添加学生信息（如果有）
    #         if analysis_result.get('knowledge', {}).get('student_id'):
    #             html_content += f"<p>学生ID: {analysis_result['knowledge']['student_id']}</p>"
            
    #         # 添加报告生成时间
    #         html_content += f"<p>报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>"
            
    #         # 添加知识点掌握情况
    #         html_content += "<div class='section'>"
    #         html_content += "<h2>知识点掌握情况</h2>"
            
    #         knowledge_mastery = analysis_result.get('knowledge', {}).get('knowledge_mastery', {})
    #         if knowledge_mastery:
    #             html_content += """
    #                 <table>
    #                     <tr>
    #                         <th>知识点</th>
    #                         <th>正确率</th>
    #                         <th>正确提交率</th>
    #                         <th>平均用时</th>
    #                         <th>总提交次数</th>
    #                         <th>正确提交次数</th>
    #                     </tr>
    #             """
                
    #             for knowledge, data in knowledge_mastery.items():
    #                 html_content += f"""
    #                     <tr>
    #                         <td>{knowledge}</td>
    #                         <td>{data['correct_rate']:.2%}</td>
    #                         <td>{data['correct_submission_rate']:.2%}</td>
    #                         <td>{data['avg_time_consume']:.2f}秒</td>
    #                         <td>{data['total_submissions']}</td>
    #                         <td>{data['correct_submissions']}</td>
    #                     </tr>
    #                 """
                
    #             html_content += "</table>"
    #         else:
    #             html_content += "<p>无知识点掌握数据</p>"
            
    #         # 添加薄弱环节分析
    #         weak_points = analysis_result.get('knowledge', {}).get('weak_points', [])
    #         if weak_points:
    #             html_content += "<h3>薄弱环节分析</h3><ul>"
                
    #             for weak_point in weak_points:
    #                 text = f"知识点 '{weak_point['knowledge']}' "
    #                 if 'sub_knowledge' in weak_point:
    #                     text += f"的从属知识点 '{weak_point['sub_knowledge']}' "
    #                 text += f"正确率为 {weak_point['correct_rate']:.2%}，{weak_point['reason']}。"
    #                 html_content += f"<li class='weak-point'>{text}</li>"
                
    #             html_content += "</ul>"
    #         html_content += "</div>"
            
    #         # 添加学习行为分析
    #         html_content += "<div class='section'>"
    #         html_content += "<h2>学习行为分析</h2>"
            
    #         behavior_data = analysis_result.get('behavior', {}).get('behavior_profile', {})
    #         if behavior_data:
    #             # 添加答题时间分布
    #             if 'peak_hours' in behavior_data:
    #                 html_content += "<h3>答题时间分布</h3>"
    #                 html_content += """
    #                     <table>
    #                         <tr>
    #                             <th>时间段</th>
    #                             <th>提交次数</th>
    #                         </tr>
    #                 """
                    
    #                 for item in behavior_data['peak_hours']:
    #                     hour = item['hour']
    #                     count = item['count']
    #                     html_content += f"""
    #                         <tr>
    #                             <td>{hour}:00-{int(hour)+1}:00</td>
    #                             <td>{count}</td>
    #                         </tr>
    #                     """
                    
    #                 html_content += "</table>"
                
    #             # 添加答题状态分布
    #             if 'state_distribution' in behavior_data:
    #                 html_content += "<h3>答题状态分布</h3>"
    #                 html_content += """
    #                     <table>
    #                         <tr>
    #                             <th>提交状态</th>
    #                             <th>次数</th>
    #                             <th>百分比</th>
    #                         </tr>
    #                 """
                    
    #                 for state, count in behavior_data['state_distribution'].items():
    #                     total = sum(behavior_data['state_distribution'].values())
    #                     percentage = count / total if total > 0 else 0
    #                     html_content += f"""
    #                         <tr>
    #                             <td>{state}</td>
    #                             <td>{count}</td>
    #                             <td>{percentage:.2%}</td>
    #                         </tr>
    #                     """
                    
    #                 html_content += "</table>"
                
    #             # 添加学习行为模式分析
    #             html_content += "<h3>学习行为模式分析</h3><ul>"
                
    #             # 创建学习行为模式文本
    #             behavior_patterns = []
                
    #             # 根据正确率分析学习效率
    #             if 'correct_rate' in behavior_data:
    #                 correct_rate = behavior_data['correct_rate']
    #                 if correct_rate > 0.8:
    #                     behavior_patterns.append("学生答题正确率很高，掌握知识点较好。")
    #                 elif correct_rate > 0.6:
    #                     behavior_patterns.append("学生答题正确率一般，对部分知识点的理解需要加强。")
    #                 else:
    #                     behavior_patterns.append("学生答题正确率较低，需要加强基础知识学习。")
                
    #             # 根据答题时间分析学习时间习惯
    #             if 'peak_hours' in behavior_data and behavior_data['peak_hours']:
    #                 peak_hour = behavior_data['peak_hours'][0]['hour']
    #                 if 8 <= peak_hour <= 12:
    #                     behavior_patterns.append("学生倾向于在上午时间段答题，精力较为充沛。")
    #                 elif 13 <= peak_hour <= 17:
    #                     behavior_patterns.append("学生倾向于在下午时间段答题，学习效率较为稳定。")
    #                 elif 18 <= peak_hour <= 22:
    #                     behavior_patterns.append("学生倾向于在晚上时间段答题，可能会影响休息时间。")
    #                 else:
    #                     behavior_patterns.append(f"学生在{peak_hour}:00-{peak_hour+1}:00时段答题最多，可能是非常规学习时间。")
                
    #             # 添加方法使用分析
    #             if 'method_distribution' in behavior_data:
    #                 methods = behavior_data['method_distribution']
    #                 if methods and len(methods) > 0:
    #                     most_method = max(methods.items(), key=lambda x: x[1])[0]
    #                     behavior_patterns.append(f"学生最常使用的答题方法是{most_method}。")
                
    #             # 如果有相对表现数据，添加相对表现分析
    #             if 'relative_performance' in behavior_data:
    #                 rel_perf = behavior_data['relative_performance']
    #                 if 'correct_rate_vs_avg' in rel_perf:
    #                     rate_vs_avg = rel_perf['correct_rate_vs_avg']
    #                     if rate_vs_avg > 0.1:
    #                         behavior_patterns.append("学生的正确率明显高于平均水平，表现优秀。")
    #                     elif rate_vs_avg < -0.1:
    #                         behavior_patterns.append("学生的正确率低于平均水平，需要改进。")
                
    #             # 添加行为模式文本
    #             if behavior_patterns:
    #                 for pattern in behavior_patterns:
    #                     html_content += f"<li>{pattern}</li>"
    #             else:
    #                 html_content += "<li>无法生成行为模式分析，数据不足。</li>"
                
    #             html_content += "</ul>"
    #         else:
    #             html_content += "<p>无学习行为分析数据</p>"
    #         html_content += "</div>"
            
    #         # 添加题目难度分析
    #         html_content += "<div class='section'>"
    #         html_content += "<h2>题目难度分析</h2>"
            
    #         difficulty_data = analysis_result.get('difficulty', {})
            
    #         # 从题目难度数据中提取评估信息
    #         question_difficulty = difficulty_data.get('question_difficulty', {})
    #         if question_difficulty:
    #             html_content += "<h3>题目难度分布</h3>"
    #             html_content += """
    #                 <table>
    #                     <tr>
    #                         <th>难度级别</th>
    #                         <th>题目数量</th>
    #                         <th>百分比</th>
    #                     </tr>
    #             """
                
    #             # 根据正确率对题目进行分类
    #             difficulty_levels = {'简单': 0, '中等': 0, '困难': 0}
                
    #             for title_id, data in question_difficulty.items():
    #                 if data['correct_rate'] > 0.7:
    #                     difficulty_levels['简单'] += 1
    #                 elif data['correct_rate'] > 0.4:
    #                     difficulty_levels['中等'] += 1
    #                 else:
    #                     difficulty_levels['困难'] += 1
                
    #             total = sum(difficulty_levels.values())
    #             for level, count in difficulty_levels.items():
    #                 percentage = count / total if total > 0 else 0
    #                 html_content += f"""
    #                     <tr>
    #                         <td>{level}</td>
    #                         <td>{count}</td>
    #                         <td>{percentage:.2%}</td>
    #                     </tr>
    #                 """
                
    #             html_content += "</table>"
            
    #             # 添加难度异常题目分析
    #             unreasonable_questions = difficulty_data.get('unreasonable_questions', [])
    #             if unreasonable_questions:
    #                 html_content += "<h3>难度异常题目分析</h3>"
    #                 html_content += """
    #                     <table>
    #                         <tr>
    #                             <th>题目ID</th>
    #                             <th>正确率</th>
    #                             <th>平均掌握度</th>
    #                             <th>知识点</th>
    #                             <th>原因</th>
    #                         </tr>
    #                 """
                    
    #                 for question in unreasonable_questions:
    #                     html_content += f"""
    #                         <tr>
    #                             <td>{question['title_id']}</td>
    #                             <td>{question['correct_rate']:.2%}</td>
    #                             <td>{question['avg_mastery']:.2%}</td>
    #                             <td>{question['knowledge']}</td>
    #                             <td>{question['reason']}</td>
    #                         </tr>
    #                     """
                    
    #                 html_content += "</table>"
    #             else:
    #                 html_content += "<p>无难度异常题目</p>"
            
    #             # 添加难度建议
    #             html_content += "<h3>难度建议</h3>"
    #             html_content += "<ul>"
                
    #             # 生成难度建议
    #             difficulty_suggestions = []
                
    #             # 根据难度分布生成建议
    #             if difficulty_levels:
    #                 # 分析难度分布情况
    #                 total_questions = sum(difficulty_levels.values())
    #                 difficult_count = difficulty_levels.get('困难', 0)
    #                 easy_count = difficulty_levels.get('简单', 0)
                    
    #                 # 如果困难题目比例过高
    #                 if difficult_count / total_questions > 0.4:
    #                     difficulty_suggestions.append("困难题目比例较高，建议适当增加中等难度和简单题目，以提高学生学习积极性。")
                    
    #                 # 如果简单题目比例过高
    #                 if easy_count / total_questions > 0.6:
    #                     difficulty_suggestions.append("简单题目比例较高，建议适当增加难度，以提升学生解决问题的能力。")
                    
    #                 # 如果难度分布较为均衡
    #                 if 0.2 <= difficult_count / total_questions <= 0.4 and 0.3 <= easy_count / total_questions <= 0.5:
    #                     difficulty_suggestions.append("题目难度分布较为均衡，适合大多数学生的学习需求。")
                
    #             # 根据异常题目生成建议
    #             if unreasonable_questions:
    #                 unreasonable_count = len(unreasonable_questions)
    #                 difficulty_suggestions.append(f"有{unreasonable_count}道题目的难度与学生知识掌握程度不符，建议检查这些题目的设计。")
                    
    #                 # 如果异常题目较多，给出进一步建议
    #                 if unreasonable_count >= 3:
    #                     difficulty_suggestions.append("建议对难度异常题目进行重新设计或提供更详细的解题指导。")
                
    #             # 添加建议到报告中
    #             if difficulty_suggestions:
    #                 for suggestion in difficulty_suggestions:
    #                     html_content += f"<li>{suggestion}</li>"
    #             else:
    #                 html_content += "<li>无难度建议</li>"
                
    #             html_content += "</ul>"
                
    #             # 此部分内容已在上面实现，不需要再次添加
    #             # html_content += "<h3>难度建议</h3><ul>"
    #             # for suggestion in difficulty_data['difficulty_suggestions']:
    #             #    html_content += f"<li>{suggestion}</li>"
    #             # html_content += "</ul>"
    #         else:
    #             html_content += "<p>无题目难度分析数据</p>"
    #         html_content += "</div>"
            
    #         # 结束HTML
    #         html_content += """
    #         </body>
    #         </html>
    #         """
            
    #         # 写入文件
    #         with open(file_path, 'w', encoding='utf-8') as f:
    #             f.write(html_content)
            
    #         print(f"已生成HTML报告: {file_path}")
            
    #     except Exception as e:
    #         print(f"生成HTML报告失败: {e}")
    #         return {'error': f'生成HTML报告失败: {e}'}

    # 包含复用
    def _generate_html_general_report(self, analysis_result, file_path, content=None):
        """生成HTML格式的综合分析报告（复用版）"""
        try:
            # 获取各部分数据
            knowledge_data = analysis_result.get('knowledge', {})
            behavior_data = analysis_result.get('behavior', {})
            difficulty_data = analysis_result.get('difficulty', {})
            
            # 获取学生ID（优先从knowledge部分获取）
            student_id = (knowledge_data.get('student_id') or 
                        behavior_data.get('student_id') or 
                        difficulty_data.get('student_id'))
            
            # 构建HTML
            html = self._generate_html_header("综合分析报告", student_id)
            
            # 添加知识点部分
            html += self._generate_knowledge_section(knowledge_data)
            
            # 添加行为分析部分
            html += self._generate_behavior_section(behavior_data)
            
            # 添加难度分析部分
            html += self._generate_difficulty_section(difficulty_data)
            
            # 结束HTML
            html += """
            </body>
            </html>
            """
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html)
            
            print(f"已生成HTML报告: {file_path}")
            
        except Exception as e:
            print(f"生成HTML报告失败: {e}")
            return {'error': f'生成HTML报告失败: {e}'}
    
    # DOCX报告生成方法
    def _generate_docx_knowledge_report(self, analysis_result, file_path, content=None):
        """生成DOCX格式的知识点掌握报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading('知识点掌握度分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加学生信息（如果有）
            if analysis_result.get('student_id'):
                p = doc.add_paragraph(f"学生ID: {analysis_result['student_id']}")
            
            # 添加报告生成时间
            doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # 添加知识点掌握情况
            doc.add_heading('知识点掌握情况', level=1)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # 设置表头
            header_cells = table.rows[0].cells
            header_cells[0].text = '知识点'
            header_cells[1].text = '正确率'
            header_cells[2].text = '正确提交率'
            header_cells[3].text = '平均用时'
            header_cells[4].text = '总提交次数'
            header_cells[5].text = '正确提交次数'
            
            # 添加数据行
            for knowledge, data in analysis_result['knowledge_mastery'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = knowledge
                row_cells[1].text = f"{data['correct_rate']:.2%}"
                row_cells[2].text = f"{data['correct_submission_rate']:.2%}"
                row_cells[3].text = f"{data['avg_time_consume']:.2f}秒"
                row_cells[4].text = str(data['total_submissions'])
                row_cells[5].text = str(data['correct_submissions'])
            
            # 添加薄弱环节分析
            if analysis_result['weak_points']:
                doc.add_heading('薄弱环节分析', level=1)
                
                for weak_point in analysis_result['weak_points']:
                    text = f"知识点 '{weak_point['knowledge']}' "
                    if 'sub_knowledge' in weak_point:
                        text += f"的从属知识点 '{weak_point['sub_knowledge']}' "
                    text += f"正确率为 {weak_point['correct_rate']:.2%}，{weak_point['reason']}。"
                    doc.add_paragraph(text)
            
            # 保存文档
            doc.save(file_path)
            print(f"已生成DOCX报告: {file_path}")
            
        except Exception as e:
            print(f"生成DOCX报告失败: {e}")
            return {'error': f'生成DOCX报告失败: {e}'}
    
    def _generate_docx_behavior_report(self, analysis_result, file_path, content=None):
        """生成DOCX格式的学习行为报告"""
        # 实现类似于_generate_docx_knowledge_report的方法
        pass
    
    def _generate_docx_difficulty_report(self, analysis_result, file_path, content=None):
        """生成DOCX格式的题目难度报告"""
        # 实现类似于_generate_docx_knowledge_report的方法
        pass
    
    def _generate_docx_general_report(self, analysis_result, file_path, content=None):
        """生成DOCX格式的综合分析报告"""
        # 实现类似于_generate_docx_knowledge_report的方法
        pass
    
    # XLSX报告生成方法
    def _generate_xlsx_knowledge_report(self, analysis_result, file_path, content=None):
        """生成XLSX格式的知识点掌握报告"""
        try:
            # 创建Excel工作簿
            writer = pd.ExcelWriter(file_path, engine='xlsxwriter')
            
            # 创建知识点掌握情况数据框
            knowledge_data = []
            for knowledge, data in analysis_result['knowledge_mastery'].items():
                knowledge_data.append({
                    '知识点': knowledge,
                    '正确率': data['correct_rate'],
                    '正确提交率': data['correct_submission_rate'],
                    '平均用时(秒)': data['avg_time_consume'],
                    '总提交次数': data['total_submissions'],
                    '正确提交次数': data['correct_submissions'],
                    '总分': data['total_score'],
                    '获得分数': data['earned_score']
                })
            
            # 转换为DataFrame并写入Excel
            knowledge_df = pd.DataFrame(knowledge_data)
            knowledge_df.to_excel(writer, sheet_name='知识点掌握情况', index=False)
            
            # 创建薄弱环节数据框
            if analysis_result['weak_points']:
                weak_points_data = []
                for weak_point in analysis_result['weak_points']:
                    data = {
                        '知识点': weak_point['knowledge'],
                        '正确率': weak_point['correct_rate'],
                        '原因': weak_point['reason']
                    }
                    if 'sub_knowledge' in weak_point:
                        data['从属知识点'] = weak_point['sub_knowledge']
                    weak_points_data.append(data)
                
                weak_points_df = pd.DataFrame(weak_points_data)
                weak_points_df.to_excel(writer, sheet_name='薄弱环节分析', index=False)
            
            # 保存Excel文件
            writer.close()
            print(f"已生成XLSX报告: {file_path}")
            
        except Exception as e:
            print(f"生成XLSX报告失败: {e}")
            return {'error': f'生成XLSX报告失败: {e}'}
    
    def _generate_xlsx_behavior_report(self, analysis_result, file_path, content=None):
        """生成XLSX格式的学习行为报告"""
        # 实现类似于_generate_xlsx_knowledge_report的方法
        pass
    
    def _generate_xlsx_difficulty_report(self, analysis_result, file_path, content=None):
        """生成XLSX格式的题目难度报告"""
        # 实现类似于_generate_xlsx_knowledge_report的方法
        pass
    
    def _generate_xlsx_general_report(self, analysis_result, file_path, content=None):
        """生成XLSX格式的综合分析报告"""
        # 实现类似于_generate_xlsx_knowledge_report的方法
        pass
        
    def get_report_history(self):
        """获取报告生成历史记录
        
        Returns:
            报告历史记录列表，每条记录包含报告ID、类型、生成时间、文件路径等信息
        """
        try:
            # 获取报告目录中的所有文件
            report_files = list(self.report_dir.glob('*.*'))
            
            # 解析文件名，提取报告信息
            reports = []
            for file_path in report_files:
                filename = file_path.name
                # 解析文件名格式: {report_type}_report[_{student_id}]_{timestamp}.{format}
                match = re.match(r'(\w+)_report(?:_(\w+))?_(\d{14})\.(\w+)', filename)
                
                if match:
                    report_type, student_id, timestamp, format_type = match.groups()
                    # 将时间戳转换为可读格式
                    try:
                        time_obj = datetime.strptime(timestamp, '%Y%m%d%H%M%S')
                        created_time = time_obj.strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        created_time = '未知时间'
                    
                    # 创建报告记录
                    report = {
                        'id': f"{report_type}_{timestamp}",
                        'type': report_type,
                        'student_id': student_id if student_id else 'all',
                        'format': format_type,
                        'created_at': created_time,
                        'file_path': str(file_path),
                        'filename': filename
                    }
                    reports.append(report)
            
            # 按时间倒序排序
            reports.sort(key=lambda x: x['created_at'], reverse=True)
            return reports
            
        except Exception as e:
            print(f"获取报告历史记录失败: {e}")
            return []
    
    def delete_report(self, report_id):
        """删除指定ID的报告
        
        Args:
            report_id: 报告ID，格式为 {report_type}_{timestamp}
            
        Returns:
            是否删除成功
        """
        try:
            # 解析报告ID，提取报告类型和时间戳
            match = re.match(r'(\w+)_(\d{14})', report_id)
            if not match:
                return False
                
            report_type, timestamp = match.groups()
            
            # 查找匹配的报告文件
            pattern = f"{report_type}_report*_{timestamp}.*"
            matching_files = list(self.report_dir.glob(pattern))
            
            if not matching_files:
                return False
                
            # 删除找到的文件
            for file_path in matching_files:
                os.remove(file_path)
                print(f"已删除报告文件: {file_path}")
                
            return True
            
        except Exception as e:
            print(f"删除报告失败: {e}")
            return False
